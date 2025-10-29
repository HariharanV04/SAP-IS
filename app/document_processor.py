#!/usr/bin/env python3
"""
Document Processor for Direct iFlow Generation
Handles various document formats (Word, PDF, text, chat) and extracts content for iFlow generation
"""

import os
import logging
import json
from pathlib import Path
import mimetypes

class DocumentProcessor:
    """Process various document formats for iFlow generation"""
    
    def __init__(self):
        """Initialize document processor"""
        self.supported_formats = {
            'text/plain': self._process_text,
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/msword': self._process_doc,
            'application/json': self._process_json,
            'text/markdown': self._process_markdown
        }
        
        # Try to import optional dependencies
        self.has_pdf_support = self._check_pdf_support()
        self.has_docx_support = self._check_docx_support()
    
    def _check_pdf_support(self):
        """Check if PDF processing is available"""
        try:
            import PyPDF2
            return True
        except ImportError:
            try:
                import pdfplumber
                return True
            except ImportError:
                logging.warning("PDF processing not available. Install PyPDF2 or pdfplumber for PDF support.")
                return False
    
    def _check_docx_support(self):
        """Check if DOCX processing is available"""
        try:
            import docx
            return True
        except ImportError:
            logging.warning("DOCX processing not available. Install python-docx for Word document support.")
            return False
    
    def process_document(self, file_path, filename):
        """
        Process a document and extract text content
        
        Args:
            file_path (str): Path to the uploaded file
            filename (str): Original filename
            
        Returns:
            dict: Processed document information
        """
        try:
            # Determine file type
            mime_type, _ = mimetypes.guess_type(filename)
            
            # Handle special cases
            if filename.lower().endswith('.txt'):
                mime_type = 'text/plain'
            elif filename.lower().endswith('.md'):
                mime_type = 'text/markdown'
            elif filename.lower().endswith('.json'):
                mime_type = 'application/json'
            elif filename.lower().endswith('.chat'):
                mime_type = 'text/plain'  # Treat chat files as text
            
            logging.info(f"Processing document: {filename} (MIME type: {mime_type})")
            
            if mime_type not in self.supported_formats:
                return {
                    'success': False,
                    'error': f'Unsupported file format: {mime_type}',
                    'supported_formats': list(self.supported_formats.keys())
                }
            
            # Process the document
            processor = self.supported_formats[mime_type]
            result = processor(file_path, filename)
            
            if result['success']:
                result.update({
                    'filename': filename,
                    'mime_type': mime_type,
                    'file_size': os.path.getsize(file_path)
                })
            
            return result
            
        except Exception as e:
            logging.error(f"Error processing document {filename}: {str(e)}")
            return {
                'success': False,
                'error': f'Error processing document: {str(e)}'
            }
    
    def _process_text(self, file_path, filename):
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                'success': True,
                'content': content,
                'content_type': 'text',
                'word_count': len(content.split()),
                'char_count': len(content)
            }
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                
                return {
                    'success': True,
                    'content': content,
                    'content_type': 'text',
                    'word_count': len(content.split()),
                    'char_count': len(content),
                    'encoding': 'latin-1'
                }
            except Exception as e:
                return {
                    'success': False,
                    'error': f'Error reading text file: {str(e)}'
                }
    
    def _process_markdown(self, file_path, filename):
        """Process Markdown files"""
        result = self._process_text(file_path, filename)
        if result['success']:
            result['content_type'] = 'markdown'
        return result
    
    def _process_json(self, file_path, filename):
        """Process JSON files (could be existing documentation)"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # If it's already processed documentation, extract the content
            if isinstance(data, dict) and 'documentation' in data:
                content = data['documentation']
            elif isinstance(data, dict) and 'content' in data:
                content = data['content']
            else:
                content = json.dumps(data, indent=2)
            
            return {
                'success': True,
                'content': content,
                'content_type': 'json',
                'json_data': data,
                'word_count': len(content.split()),
                'char_count': len(content)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing JSON file: {str(e)}'
            }
    
    def _process_pdf(self, file_path, filename):
        """Process PDF files"""
        if not self.has_pdf_support:
            return {
                'success': False,
                'error': 'PDF processing not available. Please install PyPDF2 or pdfplumber.'
            }
        
        try:
            # Try pdfplumber first (better text extraction)
            try:
                import pdfplumber
                
                text_content = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
                
                content = '\n\n'.join(text_content)
                
            except ImportError:
                # Fallback to PyPDF2
                import PyPDF2
                
                text_content = []
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
                
                content = '\n\n'.join(text_content)
            
            return {
                'success': True,
                'content': content,
                'content_type': 'pdf',
                'page_count': len(text_content),
                'word_count': len(content.split()),
                'char_count': len(content)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing PDF file: {str(e)}'
            }
    
    def _process_docx(self, file_path, filename):
        """Process DOCX files"""
        if not self.has_docx_support:
            return {
                'success': False,
                'error': 'DOCX processing not available. Please install python-docx.'
            }
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(' | '.join(row_text))

            # Extract and analyze images
            logging.info(f"Starting image extraction from {filename}")
            images_info = self._extract_images_from_docx(doc, file_path, filename)
            logging.info(f"Image extraction completed: {images_info['count']} images found, {images_info['analyzed_count']} analyzed")

            # Combine all content
            content_parts = []
            if paragraphs:
                content_parts.append('\n\n'.join(paragraphs))
            if tables_text:
                content_parts.append('\n\nTables:\n' + '\n'.join(tables_text))
            if images_info['descriptions']:
                content_parts.append('\n\nImages and Diagrams:\n' + '\n'.join(images_info['descriptions']))

            content = '\n\n'.join(content_parts)
            
            return {
                'success': True,
                'content': content,
                'content_type': 'docx',
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'image_count': images_info['count'],
                'images_analyzed': images_info['analyzed_count'],
                'word_count': len(content.split()),
                'char_count': len(content),
                'images_info': images_info
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing DOCX file: {str(e)}'
            }
    
    def _process_doc(self, file_path, filename):
        """Process DOC files (legacy Word format)"""
        return {
            'success': False,
            'error': 'Legacy DOC format not supported. Please convert to DOCX or save as text/PDF.'
        }

    def _extract_images_from_docx(self, doc, file_path, filename):
        """Extract and analyze images from DOCX document"""
        images_info = {
            'count': 0,
            'analyzed_count': 0,
            'descriptions': [],
            'image_files': []
        }

        try:
            import os
            import tempfile

            logging.info(f"Checking document relationships for images in {filename}")

            # Create temp directory for extracted images
            temp_dir = tempfile.mkdtemp(prefix=f"docx_images_{filename}_")
            logging.info(f"Created temp directory: {temp_dir}")

            # Extract images from document relationships
            total_rels = len(doc.part.rels.values())
            logging.info(f"Found {total_rels} document relationships to check")

            for rel in doc.part.rels.values():
                logging.debug(f"Checking relationship: {rel.target_ref}")
                if "image" in rel.target_ref:
                    images_info['count'] += 1
                    logging.info(f"Found image {images_info['count']}: {rel.target_ref}")
                    try:
                        # Save image to temp file
                        image_data = rel.target_part.blob
                        image_ext = rel.target_ref.split('.')[-1]
                        image_filename = f"image_{images_info['count']}.{image_ext}"
                        image_path = os.path.join(temp_dir, image_filename)

                        logging.info(f"Saving image to: {image_path} (size: {len(image_data)} bytes)")
                        with open(image_path, 'wb') as img_file:
                            img_file.write(image_data)

                        images_info['image_files'].append(image_path)

                        # Analyze image with AI (if available)
                        logging.info(f"Starting AI analysis for image {images_info['count']}")
                        description = self._analyze_image_with_ai(image_path, f"Image {images_info['count']} from {filename}")
                        if description:
                            logging.info(f"Successfully analyzed image {images_info['count']}: {len(description)} characters")
                            images_info['descriptions'].append(f"**Image {images_info['count']}:** {description}")
                            images_info['analyzed_count'] += 1
                        else:
                            logging.warning(f"Could not analyze image {images_info['count']} - adding placeholder")
                            images_info['descriptions'].append(f"**Image {images_info['count']}:** [Image found but could not be analyzed - likely contains integration diagrams, flow charts, or technical architecture]")

                    except Exception as e:
                        logging.warning(f"Error processing image {images_info['count']}: {str(e)}")
                        images_info['descriptions'].append(f"**Image {images_info['count']}:** [Image extraction failed - {str(e)}]")

            # Clean up temp directory
            try:
                import shutil
                shutil.rmtree(temp_dir)
            except:
                pass

        except Exception as e:
            logging.error(f"Error extracting images from DOCX: {str(e)}")
            images_info['descriptions'].append(f"[Error extracting images: {str(e)}]")

        return images_info

    def _analyze_image_with_ai(self, image_path, context=""):
        """Analyze image content using AI vision capabilities"""
        try:
            logging.info(f"Starting image analysis for: {image_path}")
            # Check if we have vision capabilities available
            import base64
            import os

            # Read and encode image
            with open(image_path, 'rb') as image_file:
                image_data = base64.b64encode(image_file.read()).decode('utf-8')

            # Get file extension for MIME type
            file_ext = os.path.splitext(image_path)[1].lower()
            mime_type = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif': 'image/gif',
                '.bmp': 'image/bmp'
            }.get(file_ext, 'image/png')

            # Create vision analysis prompt
            vision_prompt = f"""Analyze this image from an integration document. {context}

Focus on identifying:
1. **Integration Flow Diagrams**: Data flow between systems, process steps, decision points
2. **Architecture Diagrams**: System components, connections, data sources/targets
3. **Data Models**: Entity relationships, field mappings, data structures
4. **Technical Specifications**: APIs, endpoints, protocols, configurations
5. **Business Process**: Workflow steps, business rules, approval processes

Provide a concise technical description focusing on integration-relevant details that would help generate an iFlow.

If this appears to be a non-technical image (photos, decorative elements), respond with "NON_TECHNICAL_IMAGE".
"""

            # Try to use Claude Sonnet-4 with vision capabilities
            try:
                logging.info("Importing DocumentationEnhancer for vision analysis")
                try:
                    from app.documentation_enhancer import DocumentationEnhancer
                except ImportError:
                    # Fallback for when running from different directory
                    from documentation_enhancer import DocumentationEnhancer
                enhancer = DocumentationEnhancer()

                logging.info(f"Calling Anthropic vision API for {mime_type} image")
                # Call Anthropic with vision
                response = enhancer.analyze_image_with_anthropic(vision_prompt, image_data, mime_type)

                if response and "NON_TECHNICAL_IMAGE" not in response:
                    logging.info(f"Vision analysis successful: {len(response)} characters")
                    return response.strip()
                else:
                    logging.info("Vision analysis returned non-technical image or empty response")
                    return None

            except Exception as e:
                logging.error(f"AI vision analysis failed: {str(e)}")
                return None

        except Exception as e:
            logging.error(f"Error analyzing image {image_path}: {str(e)}")
            return None
    
    def generate_documentation_json(self, processed_doc, job_id, llm_provider='anthropic'):
        """
        Generate intermediate JSON for iFlow generation from processed documentation

        Args:
            processed_doc (dict): Processed document information
            job_id (str): Job ID for tracking

        Returns:
            dict: Documentation JSON for iFlow generation
        """
        try:
            # Debug logging
            print(f"DEBUG: DocumentProcessor.generate_documentation_json called with LLM provider: {llm_provider}")
            logging.info(f"DocumentProcessor using LLM provider: {llm_provider}")

            # Convert raw content to structured markdown using LLM
            markdown_content = self._convert_to_markdown(processed_doc, llm_provider)

            documentation_json = {
                'job_id': job_id,
                'source_type': 'uploaded_documentation',
                'source_file': processed_doc.get('filename', 'unknown'),
                'content_type': processed_doc.get('content_type', 'text'),
                'documentation': markdown_content,  # Now using converted markdown
                'original_content': processed_doc['content'],  # Keep original for reference
                'metadata': {
                    'word_count': processed_doc.get('word_count', 0),
                    'char_count': processed_doc.get('char_count', 0),
                    'file_size': processed_doc.get('file_size', 0),
                    'mime_type': processed_doc.get('mime_type', 'unknown'),
                    'converted_to_markdown': True
                },
                'processing_info': {
                    'skip_document_generation': True,
                    'ready_for_iflow_generation': True,
                    'processing_timestamp': str(datetime.now())
                }
            }
            
            # Add format-specific metadata
            if processed_doc.get('content_type') == 'pdf':
                documentation_json['metadata']['page_count'] = processed_doc.get('page_count', 0)
            elif processed_doc.get('content_type') == 'docx':
                documentation_json['metadata']['paragraph_count'] = processed_doc.get('paragraph_count', 0)
                documentation_json['metadata']['table_count'] = processed_doc.get('table_count', 0)
            elif processed_doc.get('content_type') == 'json':
                documentation_json['json_data'] = processed_doc.get('json_data', {})
            
            return {
                'success': True,
                'documentation_json': documentation_json
            }
            
        except Exception as e:
            logging.error(f"Error generating documentation JSON: {str(e)}")
            return {
                'success': False,
                'error': f'Error generating documentation JSON: {str(e)}'
            }

    def _convert_to_markdown(self, processed_doc, llm_provider='anthropic'):
        """
        Convert raw document content to structured markdown using LLM

        Args:
            processed_doc (dict): Processed document information

        Returns:
            str: Structured markdown content suitable for iFlow generation
        """
        try:
            # Import here to avoid circular imports
            from documentation_enhancer import DocumentationEnhancer

            raw_content = processed_doc['content']
            content_type = processed_doc.get('content_type', 'text')
            filename = processed_doc.get('filename', 'unknown')

            # If content is already markdown, return as-is
            if content_type == 'markdown':
                return raw_content

            # Create documentation enhancer instance with selected LLM provider
            print(f"DEBUG: Creating DocumentationEnhancer with LLM provider: {llm_provider}")
            logging.info(f"Creating DocumentationEnhancer with LLM provider: {llm_provider}")
            enhancer = DocumentationEnhancer(selected_service=llm_provider)

            # Prepare prompt for markdown conversion
            conversion_prompt = f"""
You are an expert technical documentation analyst. Convert the following {content_type} document content into well-structured markdown format suitable for integration flow generation.

**Source Document:** {filename}
**Content Type:** {content_type}

**Instructions:**
1. Structure the content with proper markdown headers (# ## ###)
2. Identify and highlight key integration components, APIs, data flows, and business processes
3. Pay special attention to any image descriptions - they often contain crucial integration diagrams, architecture flows, and technical specifications
4. Create clear sections for:
   - Overview/Purpose
   - Integration Requirements
   - Data Sources and Destinations
   - Business Logic/Rules
   - Error Handling Requirements
   - Technical Specifications
4. Use bullet points, numbered lists, and tables where appropriate
5. Preserve all technical details and specifications
6. Format code snippets, URLs, and technical terms properly
7. Ensure the output is clean, readable markdown that clearly describes integration requirements

**Original Content:**
{raw_content}

**Output:** Return only the converted markdown content, no additional commentary.
"""

            # Call LLM to convert content using DocumentationEnhancer
            if llm_provider == 'gemma3':
                # For Gemma-3, we need to call the Gemma-3 API directly
                response = self._call_gemma3_for_conversion(conversion_prompt)
            else:
                # Use Anthropic
                response = enhancer.enhance_with_anthropic(conversion_prompt)

            if response and response.strip():
                logging.info(f"Successfully converted {content_type} content to markdown for {filename}")
                return response.strip()
            else:
                logging.warning(f"LLM conversion failed for {filename}, using original content")
                return self._fallback_markdown_conversion(raw_content, filename, content_type)

        except Exception as e:
            logging.error(f"Error converting content to markdown: {str(e)}")
            # Fallback to basic markdown conversion
            return self._fallback_markdown_conversion(
                processed_doc['content'],
                processed_doc.get('filename', 'unknown'),
                processed_doc.get('content_type', 'text')
            )

    def _fallback_markdown_conversion(self, content, filename, content_type):
        """
        Fallback method to create basic markdown structure when LLM conversion fails

        Args:
            content (str): Raw content
            filename (str): Source filename
            content_type (str): Content type

        Returns:
            str: Basic markdown formatted content
        """
        markdown_content = f"""# Integration Documentation

**Source:** {filename}
**Type:** {content_type}
**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## Overview

{content}

## Integration Requirements

*Please review the content above and identify:*
- Data sources and destinations
- Business logic requirements
- Error handling needs
- Technical specifications

---
*Note: This document was automatically converted from {content_type} format. Please review and enhance as needed.*
"""

    def _call_gemma3_for_conversion(self, prompt):
        """Call Gemma-3 API for document conversion"""
        try:
            import requests
            import os
            import time

            gemma3_api_url = os.getenv('GEMMA3_API_URL', 'http://localhost:5002')
            if gemma3_api_url.endswith('/api'):
                gemma3_api_url = gemma3_api_url[:-4]

            # Call Gemma-3 API for document conversion
            logging.info("Calling Gemma-3 API for document conversion")

            response = requests.post(
                f"{gemma3_api_url}/api/generate-iflow",
                json={
                    "markdown": prompt,
                    "iflow_name": "DocumentConversion",
                    "platform": "document_conversion"
                },
                timeout=1200  # 20 minutes for cold start
            )

            if response.status_code == 202:
                # Get job ID and poll for results
                result = response.json()
                job_id = result.get('job_id')

                if job_id:
                    # Poll for completion
                    for attempt in range(240):  # 20 minutes max
                        time.sleep(5)
                        status_response = requests.get(f"{gemma3_api_url}/api/jobs/{job_id}")
                        if status_response.status_code == 200:
                            status_data = status_response.json()
                            if status_data.get('status') == 'completed':
                                return status_data.get('final_response', prompt)
                            elif status_data.get('status') == 'failed':
                                logging.error("Gemma-3 document conversion failed")
                                break

            logging.warning("Gemma-3 API call failed, using basic conversion")
            return prompt.split("**Content:**")[-1].strip() if "**Content:**" in prompt else prompt

        except Exception as e:
            logging.error(f"Error calling Gemma-3 API for conversion: {str(e)}")
            return None

# Import datetime at the top level
from datetime import datetime
